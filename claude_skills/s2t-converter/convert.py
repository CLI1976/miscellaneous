#!/usr/bin/env python3
"""
簡體中文轉繁體中文轉換工具
支援格式: .md, .txt, .html, .docx
"""

import sys
import os
import argparse
from pathlib import Path

try:
    from opencc import OpenCC
except ImportError:
    print("錯誤: 請先安裝 opencc-python-reimplemented")
    print("執行: pip install opencc-python-reimplemented")
    sys.exit(1)


def convert_text(text: str, config: str = "s2twp") -> str:
    """
    轉換文字

    config 選項:
    - s2t: 簡體到繁體
    - s2tw: 簡體到台灣繁體
    - s2twp: 簡體到台灣繁體 (含慣用詞轉換，如「軟件」→「軟體」)
    - s2hk: 簡體到香港繁體
    """
    cc = OpenCC(config)
    return cc.convert(text)


def convert_txt_file(input_path: Path, output_path: Path, config: str = "s2twp") -> None:
    """轉換純文字檔案 (.txt, .md, .html)"""
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'big5']

    content = None
    for encoding in encodings:
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        raise ValueError(f"無法讀取檔案，嘗試的編碼: {encodings}")

    converted = convert_text(content, config)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(converted)


def convert_docx_file(input_path: Path, output_path: Path, config: str = "s2twp") -> None:
    """轉換 Word 文件 (.docx)"""
    try:
        from docx import Document
    except ImportError:
        print("錯誤: 請先安裝 python-docx")
        print("執行: pip install python-docx")
        sys.exit(1)

    doc = Document(input_path)
    cc = OpenCC(config)

    # 轉換段落
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = cc.convert(run.text)

    # 轉換表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = cc.convert(run.text)

    # 轉換頁首頁尾
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = cc.convert(run.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = cc.convert(run.text)

    doc.save(output_path)


def get_output_path(input_path: Path, suffix: str = "_tc") -> Path:
    """生成輸出檔案路徑"""
    stem = input_path.stem
    ext = input_path.suffix
    return input_path.parent / f"{stem}{suffix}{ext}"


def convert_file(input_path: str, output_path: str = None, config: str = "s2twp") -> str:
    """
    主要轉換函數

    Args:
        input_path: 輸入檔案路徑
        output_path: 輸出檔案路徑 (可選，預設加上 _tc 後綴)
        config: OpenCC 配置 (s2t, s2tw, s2twp, s2hk)

    Returns:
        輸出檔案路徑
    """
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f"找不到檔案: {input_path}")

    if output_path:
        output_path = Path(output_path)
    else:
        output_path = get_output_path(input_path)

    ext = input_path.suffix.lower()

    if ext in ['.txt', '.md', '.html', '.htm', '.xml', '.json', '.csv']:
        convert_txt_file(input_path, output_path, config)
    elif ext == '.docx':
        convert_docx_file(input_path, output_path, config)
    elif ext == '.doc':
        raise ValueError("不支援 .doc 格式，請先轉換為 .docx")
    else:
        # 嘗試當作純文字處理
        print(f"警告: 未知格式 {ext}，嘗試當作純文字處理")
        convert_txt_file(input_path, output_path, config)

    return str(output_path)


def main():
    parser = argparse.ArgumentParser(
        description='簡體中文轉繁體中文工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
轉換模式說明:
  s2t   - 簡體到繁體 (基本轉換)
  s2tw  - 簡體到台灣繁體
  s2twp - 簡體到台灣繁體 + 慣用詞 (預設，推薦)
  s2hk  - 簡體到香港繁體

範例:
  python convert.py document.md
  python convert.py document.docx -o output.docx
  python convert.py document.txt -c s2hk
        """
    )

    parser.add_argument('input', help='輸入檔案路徑')
    parser.add_argument('-o', '--output', help='輸出檔案路徑 (預設: 原檔名_tc.副檔名)')
    parser.add_argument('-c', '--config', default='s2twp',
                        choices=['s2t', 's2tw', 's2twp', 's2hk'],
                        help='轉換模式 (預設: s2twp)')

    args = parser.parse_args()

    try:
        output = convert_file(args.input, args.output, args.config)
        print(f"[OK] Convert complete: {output}")
    except Exception as e:
        print(f"[ERROR] {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
